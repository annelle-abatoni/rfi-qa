import os
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

def ensure_arial_font(doc):
    """Ensure all runs in the document use Arial font."""
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = 'Arial'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

def number_titles_from_page_4(doc):
    """Apply numbering and styling to titles from page 4 onward."""
    title_count = 1
    page_break_found = 0

    for para in doc.paragraphs:
        # Count page breaks (each run containing only a page break character)
        if any(run.text == '\f' for run in para.runs):
            page_break_found += 1

        if page_break_found >= 3:  # After 3rd page break (i.e., from page 4)
            if para.style.name.startswith('Heading') or para.text.isupper():  # crude check for titles
                original_text = para.text.strip()
                if original_text:
                    para.text = f"{title_count}. {original_text}"
                    para.style.font.name = 'Arial'
                    para.style.font.size = Pt(12)
                    para.style.font.bold = True
                    title_count += 1

def format_doc(input_path, output_path):
    if not os.path.exists(input_path):
        print(f"❌ File not found: {input_path}")
        return

    doc = Document(input_path)
    ensure_arial_font(doc)
    number_titles_from_page_4(doc)
    doc.save(output_path)
    print(f"✅ Formatted document saved at: {output_path}")

# --- Run the script ---
if __name__ == "__main__":
    from docx.oxml.ns import qn

    input_path = input("Enter the full path to your .docx file: ").strip('" ')
    output_path = os.path.splitext(input_path)[0] + "_formatted.docx"
    format_doc(input_path, output_path)