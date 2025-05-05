import os
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

def format_doc_with_custom_filename(input_file):
    if not os.path.exists(input_file):
        print(f"❌ File not found: {input_file}")
        return

    folder_path = os.path.dirname(input_file)
    folder_name = os.path.basename(folder_path)
    folder_number_match = re.search(r"\d+", folder_name)
    folder_number = folder_number_match.group() if folder_number_match else "000"

    base_file_name = os.path.splitext(os.path.basename(input_file))[0]
    base_file_name_lower = base_file_name.lower()

    if "sop" in base_file_name_lower:
        code = "SOP"
        suffix = folder_number
    elif "measurement of uncertainty" in base_file_name_lower:
        code = "PRT"
        suffix = folder_number + "A"
    elif "validation report" in base_file_name_lower:
        code = "REP"
        suffix = folder_number + "B"
    else:
        code = None

    if code:
        formatted_name = f"DRG-{code}-{suffix}-{base_file_name}-V01.0"
    else:
        formatted_name = f"DRG_{base_file_name}"

    output_file_name = formatted_name + ".docx"
    output_path = os.path.join(folder_path, output_file_name)

    doc = Document(input_file)

    for section in doc.sections:
        # --- Header formatting ---
        header = section.header
        header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_paragraph.clear()
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_paragraph.add_run(base_file_name)
        run.italic = True
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x18, 0x0A, 0x5E)

    doc.save(output_path)
    print(f"✅ Document saved at: {output_path}")

# --- Run the script ---
input_path = input("Enter the full path to your .docx file: ").strip('" ')
format_doc_with_custom_filename(input_path)
